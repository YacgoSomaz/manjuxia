import asyncio
import json
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from database.db import get_db
from models.novels import NovelCreate, NovelResponse, ChapterCreate, ChapterUpdate, ChapterResponse
from services.novel_service import NovelService
from services.novel_creation_service import NovelCreationService
from services.script_service import ScriptService
from services.log_service import check_novel_running
from services.llm_service import LLMService
from services.template_service import get_by_id as get_template_by_id
from utils.timezone import now_beijing_str

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/novels", tags=["novels"])

# ============ 内存中的任务状态追踪 ============
_generation_status: Dict[int, Dict[str, Any]] = {}
_rewrite_tasks: Dict[int, Dict[str, Any]] = {}


# ============ 请求模型 ============
class CreateOutlineRequest(BaseModel):
    novel_name: str
    concept: str
    llm_config_id: int
    template_id: Optional[int] = None


class GenerateChapterRequest(BaseModel):
    chapter_index: int
    llm_config_id: int
    template_id: Optional[int] = None


class UpdateOutlineRequest(BaseModel):
    outline: str


class UpdateWritingContextRequest(BaseModel):
    content: Optional[str] = None
    dynamic_state: Optional[str] = None


class ScriptToNovelRequest(BaseModel):
    name: str
    content: str
    template_id: int
    llm_config_id: int


class ScriptToScriptRequest(BaseModel):
    # v3.61.252: 剧本转剧本 — 导入原始剧本,按集切分后逐集格式化成标准剧本
    name: str
    content: str
    template_id: int
    llm_config_id: int


class NovelIncrementalImportRequest(BaseModel):
    raw_content: str


async def _read_upload_raw_content(file: UploadFile) -> tuple[str, str]:
    """读取 txt/docx 上传内容,返回 (filename, raw_content)。"""
    filename = file.filename or ""
    fname_lower = filename.lower()

    if not (fname_lower.endswith('.txt') or fname_lower.endswith('.docx')):
        raise HTTPException(status_code=400, detail="只支持 .txt 或 .docx 文件")

    content = await file.read()
    if fname_lower.endswith('.docx'):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            raw_content = "\n\n".join(paragraphs)
            if not raw_content.strip():
                raise HTTPException(status_code=400, detail=".docx 文件内容为空或无法解析")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f".docx 解析失败: {type(e).__name__}: {e}")
    else:
        try:
            raw_content = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                raw_content = content.decode('gbk')
            except UnicodeDecodeError:
                try:
                    raw_content = content.decode('gb2312')
                except UnicodeDecodeError:
                    raise HTTPException(status_code=400, detail="文件编码无法识别，请使用 UTF-8 或 GBK 编码")
    return filename, raw_content


class BatchRewriteRequest(BaseModel):
    chapter_ids: List[int]
    llm_config_id: int
    template_id: int


# ============ 基础 CRUD 端点 ============

@router.get("/")
async def get_novels():
    """获取小说列表"""
    return await NovelService.get_all()


@router.post("/", response_model=NovelResponse)
async def create_novel(novel: NovelCreate):
    """创建小说（JSON方式）"""
    return await NovelService.create(novel)


@router.post("/upload", response_model=NovelResponse)
async def upload_novel(
    file: UploadFile = File(...),
    mode: str = Form("import"),
    name: str = Form(""),
):
    """上传 .txt 或 .docx 文件导入小说/剧本。mode: import(小说) / script_import(剧本)"""
    filename = file.filename or ""
    fname_lower = filename.lower()

    if not (fname_lower.endswith('.txt') or fname_lower.endswith('.docx')):
        raise HTTPException(status_code=400, detail="只支持 .txt 或 .docx 文件")

    # 读取文件内容
    content = await file.read()

    # 根据扩展名分流解析
    if fname_lower.endswith('.docx'):
        # .docx 用 python-docx 抽段落文本
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            raw_content = "\n\n".join(paragraphs)
            if not raw_content.strip():
                raise HTTPException(status_code=400, detail=".docx 文件内容为空或无法解析")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f".docx 解析失败: {type(e).__name__}: {e}")
    else:
        # .txt 尝试 UTF-8 / GBK / GB2312
        try:
            raw_content = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                raw_content = content.decode('gbk')
            except UnicodeDecodeError:
                try:
                    raw_content = content.decode('gb2312')
                except UnicodeDecodeError:
                    raise HTTPException(status_code=400, detail="文件编码无法识别，请使用 UTF-8 或 GBK 编码")

    # 优先使用用户填写的名称;留空则用文件名(去掉扩展名)
    name = (name or "").strip() or filename.rsplit('.', 1)[0]

    novel = NovelCreate(name=name, raw_content=raw_content, mode=(mode or "import"))
    return await NovelService.create(novel)


@router.get("/{novel_id}", response_model=NovelResponse)
async def get_novel(novel_id: int):
    """获取小说详情"""
    novel = await NovelService.get_by_id(novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="小说不存在")
    return novel


@router.get("/{novel_id}/template-context")
async def get_novel_template_context(novel_id: int):
    """返回当前小说各阶段已使用的模板ID,用于前端按小说动态推荐下游模板

    返回示例:
      { "novel_id": 33,
        "script_to_novel_template_id": 85,    # 当初剧本逆转小说用的模板 (如果是该路径创建)
        "script_conversion_template_id": 80,  # 最近一次小说转剧本用的模板
        "storyboard_template_id": 25         # 最近一次分镜生成用的模板
      }
    字段为 null 表示该阶段还未生成过
    """
    from database.db import get_db
    db = await get_db()
    try:
        # 小说自身 template_id (仅 script_to_novel 模式下有值)
        cur = await db.execute("SELECT template_id FROM novels WHERE id = ?", (novel_id,))
        row = await cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="小说不存在")
        novel_template_id = row["template_id"] if "template_id" in row.keys() else None

        # 最近的 scripts.template_id (按 created_at DESC),同时带回 genres 给前端推荐用。
        # 前端只需要 genres 做推荐,不应再为推荐额外 getTemplate(template_id) 触发模板内容拉取/时序问题。
        cur = await db.execute(
            """
            SELECT s.template_id, pt.genres
            FROM scripts s
            LEFT JOIN prompt_templates pt ON pt.id = s.template_id
            WHERE s.novel_id = ? AND s.template_id IS NOT NULL
            ORDER BY s.created_at DESC
            LIMIT 1
            """,
            (novel_id,)
        )
        r2 = await cur.fetchone()
        script_template_id = r2["template_id"] if r2 else None
        script_template_genres = r2["genres"] if r2 and "genres" in r2.keys() else None

        # 最近的 storyboards.template_id
        cur = await db.execute(
            "SELECT template_id FROM storyboards WHERE novel_id = ? AND template_id IS NOT NULL ORDER BY created_at DESC LIMIT 1",
            (novel_id,)
        )
        r3 = await cur.fetchone()
        storyboard_template_id = r3["template_id"] if r3 else None

        return {
            "novel_id": novel_id,
            "script_to_novel_template_id": novel_template_id,
            "script_conversion_template_id": script_template_id,
            "script_conversion_template_genres": script_template_genres,
            "storyboard_template_id": storyboard_template_id,
        }
    finally:
        await db.close()


@router.delete("/{novel_id}")
async def delete_novel(novel_id: int, force: bool = False):
    """删除小说(级联清理 DB + 磁盘 + 终止队列任务)

    - force=False(默认): 队列里有 generating 任务时返回 409
    - force=True: 强制终止后删除
    """
    result = await NovelService.delete(novel_id, force=force)
    if not result:
        raise HTTPException(status_code=404, detail="小说不存在")
    return {
        "message": "删除成功",
        "deleted_files": result.get("deleted_files", 0) if isinstance(result, dict) else 0,
        "freed_bytes": result.get("freed_bytes", 0) if isinstance(result, dict) else 0,
    }


@router.post("/{novel_id}/parse-chapters")
async def parse_chapters(novel_id: int):
    """自动解析章节"""
    # 检查小说是否存在
    novel = await NovelService.get_by_id(novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="小说不存在")
    
    count = await NovelService.parse_chapters(novel_id)
    return {"message": f"成功解析 {count} 个章节", "chapter_count": count}


@router.post("/{novel_id}/incremental-import")
async def incremental_import_chapters(novel_id: int, request: NovelIncrementalImportRequest):
    """增量导入章节:同集号/章节号更新,不存在则新增;不清空旧章节。"""
    if not request.raw_content or not request.raw_content.strip():
        raise HTTPException(status_code=400, detail="导入内容不能为空")

    result = await NovelService.incremental_import_chapters(novel_id, request.raw_content)
    if result is None:
        raise HTTPException(status_code=404, detail="小说不存在")
    if result["total"] <= 0:
        raise HTTPException(status_code=400, detail="未识别到可导入的章节")
    return {"message": f"增量导入完成: 更新 {result['updated']} 集,新增 {result['created']} 集", **result}


@router.post("/{novel_id}/incremental-upload")
async def incremental_upload_chapters(novel_id: int, file: UploadFile = File(...)):
    """上传 .txt/.docx 增量导入到已有小说。"""
    _filename, raw_content = await _read_upload_raw_content(file)
    result = await NovelService.incremental_import_chapters(novel_id, raw_content)
    if result is None:
        raise HTTPException(status_code=404, detail="小说不存在")
    if result["total"] <= 0:
        raise HTTPException(status_code=400, detail="未识别到可导入的章节")
    return {"message": f"增量导入完成: 更新 {result['updated']} 集,新增 {result['created']} 集", **result}


@router.get("/{novel_id}/chapters", response_model=List[ChapterResponse])
async def get_chapters(novel_id: int):
    """获取章节列表"""
    # 检查小说是否存在
    novel = await NovelService.get_by_id(novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="小说不存在")
    
    return await NovelService.get_chapters(novel_id)


@router.post("/{novel_id}/chapters", response_model=ChapterResponse)
async def add_chapter(novel_id: int, chapter: ChapterCreate):
    """手动添加章节"""
    # 检查小说是否存在
    novel = await NovelService.get_by_id(novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="小说不存在")
    
    result = await NovelService.add_chapter(novel_id, chapter)
    return result


@router.put("/chapters/{chapter_id}", response_model=ChapterResponse)
async def update_chapter(chapter_id: int, chapter: ChapterUpdate):
    """更新章节"""
    result = await NovelService.update_chapter(chapter_id, chapter)
    if not result:
        raise HTTPException(status_code=404, detail="章节不存在")
    return result


@router.delete("/chapters/{chapter_id}")
async def delete_chapter(chapter_id: int):
    """删除章节"""
    result = await NovelService.delete_chapter(chapter_id)
    if not result:
        raise HTTPException(status_code=404, detail="章节不存在")
    return {"message": "删除成功"}


@router.put("/{novel_id}/chapters/reorder")
async def reorder_chapters(novel_id: int, chapter_ids: List[int]):
    """重新排序章节"""
    # 检查小说是否存在
    novel = await NovelService.get_by_id(novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="小说不存在")
    
    result = await NovelService.reorder_chapters(novel_id, chapter_ids)
    return {"message": "排序成功"}


# ============ 创作模式端点 ============

@router.post("/create-outline")
async def create_outline(request: CreateOutlineRequest):
    """生成小说大纲（创作模式）"""
    db = await get_db()
    try:
        # 创建小说记录
        now = now_beijing_str()
        await NovelService.ensure_owner_columns(db)
        owner_user_id, owner_team_id, owner_seat_id = NovelService.current_owner_values()
        cursor = await db.execute(
            """INSERT INTO novels
                   (name, raw_content, mode, created_at, updated_at, owner_user_id, owner_team_id, owner_seat_id)
               VALUES (?, '', 'create', ?, ?, ?, ?, ?)""",
            (request.novel_name, now, now, owner_user_id, owner_team_id, owner_seat_id)
        )
        await db.commit()
        novel_id = cursor.lastrowid
    finally:
        await db.close()

    # 生成大纲
    try:
        outline = await NovelCreationService.generate_outline(
            novel_id=novel_id,
            concept=request.concept,
            config_id=request.llm_config_id,
            template_id=request.template_id
        )
        return {"novel_id": novel_id, "outline": outline}
    except Exception as e:
        logger.error(f"生成大纲失败: {e}")
        raise HTTPException(status_code=500, detail=f"生成大纲失败: {str(e)}")


@router.put("/{novel_id}/outline")
async def update_outline(novel_id: int, request: UpdateOutlineRequest):
    """更新小说大纲"""
    db = await get_db()
    try:
        await db.execute(
            "UPDATE novels SET outline = ? WHERE id = ?",
            (request.outline, novel_id)
        )
        await db.commit()
        return {"message": "大纲已更新"}
    finally:
        await db.close()


@router.post("/{novel_id}/generate-chapter")
async def generate_chapter(novel_id: int, request: GenerateChapterRequest):
    """生成单个章节（异步提交）"""
    # 检查是否有正在进行的生成任务
    if novel_id in _generation_status:
        status = _generation_status[novel_id]
        if status.get("status") in ("generating", "post_processing"):
            raise HTTPException(status_code=409, detail="当前有章节正在生成中，请等待完成")

    # 标记为生成中
    _generation_status[novel_id] = {"status": "generating", "chapter_index": request.chapter_index}

    async def _do_generate():
        try:
            chapter_result = await NovelCreationService.generate_chapter(
                novel_id=novel_id,
                chapter_index=request.chapter_index,
                config_id=request.llm_config_id,
                template_id=request.template_id
            )
            _generation_status[novel_id] = {"status": "post_processing", "chapter_index": request.chapter_index}

            # 后处理
            await NovelCreationService.post_chapter_process(
                novel_id=novel_id,
                chapter_id=chapter_result["chapter_id"],
                config_id=request.llm_config_id
            )
            _generation_status[novel_id] = {"status": "success", "chapter_index": request.chapter_index}
        except Exception as e:
            logger.error(f"章节生成失败: {e}")
            _generation_status[novel_id] = {"status": "error", "error": str(e), "chapter_index": request.chapter_index}

    asyncio.create_task(_do_generate())
    return {"message": "章节生成任务已提交"}


@router.get("/{novel_id}/generation-status")
async def get_generation_status(novel_id: int):
    """获取章节生成状态"""
    # 查询最近的相关日志
    db = await get_db()
    try:
        cursor = await db.execute(
            """SELECT id, task_type, status, chapter_title, error_message,
               start_time, end_time, output_content
            FROM llm_logs
            WHERE novel_id = ? AND task_type IN ('novel_creation', 'novel_post_process')
            ORDER BY id DESC LIMIT 5""",
            (novel_id,)
        )
        logs = [dict(row) for row in await cursor.fetchall()]
    finally:
        await db.close()

    current_log = logs[0] if logs else None

    # 也检查内存中的状态（异步任务可能还未写入日志）
    mem_status = _generation_status.get(novel_id, {})

    # 判断状态：优先使用内存状态，回退到日志状态
    if mem_status.get("status") in ("generating", "post_processing"):
        status = mem_status["status"]
    elif mem_status.get("status") == "error":
        status = "failed"
    elif mem_status.get("status") == "success":
        status = "success"
    elif current_log and current_log["status"] == "running":
        if current_log["task_type"] == "novel_creation":
            status = "generating"
        else:
            status = "post_processing"
    elif current_log and current_log["status"] == "error":
        status = "failed"
    elif current_log and current_log["status"] == "success":
        status = "success"
    else:
        status = "idle"

    # 如果内存中有错误信息，注入到 current_log
    if current_log is None and mem_status.get("error"):
        current_log = {"error_message": mem_status["error"]}
    elif current_log and mem_status.get("error") and not current_log.get("error_message"):
        current_log["error_message"] = mem_status["error"]

    # 计算耗时
    if current_log and current_log.get("start_time") and current_log.get("end_time"):
        try:
            from datetime import datetime
            start = datetime.fromisoformat(str(current_log["start_time"]))
            end = datetime.fromisoformat(str(current_log["end_time"]))
            current_log["duration_seconds"] = (end - start).total_seconds()
        except Exception:
            pass

    return {
        "status": status,
        "current_log": current_log,
        "recent_logs": logs
    }


@router.get("/{novel_id}/writing-context")
async def get_writing_context(novel_id: int):
    """获取创作上下文"""
    try:
        return await NovelCreationService.get_writing_context(novel_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/writing-context/{context_id}")
async def update_writing_context(context_id: int, request: UpdateWritingContextRequest):
    """更新创作上下文"""
    try:
        result = await NovelCreationService.update_writing_context(
            context_id=context_id,
            content=request.content,
            dynamic_state=request.dynamic_state
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============ 剧本转小说端点 ============

@router.post("/script-to-novel")
async def script_to_novel(request: ScriptToNovelRequest):
    """剧本转小说"""
    # 先创建小说以获取 novel_id，再做冲突检查
    db = await get_db()
    try:
        # 创建小说记录
        now = now_beijing_str()
        await NovelService.ensure_owner_columns(db)
        owner_user_id, owner_team_id, owner_seat_id = NovelService.current_owner_values()
        cursor = await db.execute(
            """INSERT INTO novels
                   (name, raw_content, mode, template_id, created_at, updated_at, owner_user_id, owner_team_id, owner_seat_id)
               VALUES (?, '', 'script_to_novel', ?, ?, ?, ?, ?, ?)""",
            (request.name, request.template_id, now, now, owner_user_id, owner_team_id, owner_seat_id)
        )
        await db.commit()
        novel_id = cursor.lastrowid
    finally:
        await db.close()

    # 冲突检查（新创建的小说理论上不会有冲突，但保持一致性）
    db = await get_db()
    try:
        await check_novel_running(novel_id, db)
    finally:
        await db.close()

    try:
        # 获取模板
        template = await get_template_by_id(request.template_id)
        if not template:
            raise HTTPException(status_code=404, detail="模板不存在")

        template_content = template["content"]
        # 渲染模板 - 兼容多种变量名
        prompt = template_content
        replaced = False
        content_to_replace = request.content
        variable_names = ["content", "novel_content", "chapter_content", "text", "script_content"]
        for var_name in variable_names:
            for fmt in ["{" + var_name + "}", "{{" + var_name + "}}"]:
                if fmt in prompt:
                    prompt = prompt.replace(fmt, content_to_replace)
                    replaced = True
        # 无占位符回退：追加到模板末尾
        if not replaced:
            prompt = template_content + "\n\n" + content_to_replace

        messages = [
            {"role": "system", "content": "你是一位专业的小说作家，擅长将剧本改编为小说。请根据提供的剧本内容，创作出生动的小说文本。"},
            {"role": "user", "content": prompt}
        ]

        logger.info(f"开始剧本转小说: novel_id={novel_id}, name={request.name}")
        result = await LLMService.call_llm(
            config_id=request.llm_config_id,
            messages=messages,
            timeout=600,
            max_tokens=16384,
            task_type="script_to_novel",
            novel_id=novel_id
        )

        if not result or not result.strip():
            raise ValueError("大模型返回空内容")

        # 保存为章节
        db = await get_db()
        try:
            # 更新小说的 raw_content
            await db.execute(
                "UPDATE novels SET raw_content = ?, updated_at = ? WHERE id = ?",
                (result, now_beijing_str(), novel_id)
            )
            await db.commit()
        finally:
            await db.close()

        # 自动解析章节
        chapter_count = await NovelService.parse_chapters(novel_id)

        return {
            "novel_id": novel_id,
            "message": f"剧本转小说成功，已拆分为 {chapter_count} 个章节",
            "chapter_count": chapter_count
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"剧本转小说失败: {e}")
        raise HTTPException(status_code=500, detail=f"剧本转小说失败: {str(e)}")


# ============ 剧本转剧本端点 ============

@router.post("/script-to-script")
async def script_to_script(request: ScriptToScriptRequest):
    """v3.61.252: 剧本转剧本

    跟「剧本转小说」根本区别(后者是整篇 LLM 改写成小说 → 再猜集边界,会丢集/截断):
    - raw_content 直接存用户导入的原始剧本(不经 LLM)
    - 集边界由 parse_chapters(同漫剧切章正则)代码定死 → 集数永不错位、不截断
    - 逐集复用 ScriptService 章节→剧本格式化,写 scripts,失败明细原样返回不吞
    """
    # 1. 建小说(mode=script_to_script,raw_content=原始剧本)
    db = await get_db()
    try:
        now = now_beijing_str()
        await NovelService.ensure_owner_columns(db)
        owner_user_id, owner_team_id, owner_seat_id = NovelService.current_owner_values()
        cursor = await db.execute(
            """INSERT INTO novels
                   (name, raw_content, mode, template_id, created_at, updated_at, owner_user_id, owner_team_id, owner_seat_id)
               VALUES (?, ?, 'script_to_script', ?, ?, ?, ?, ?, ?)""",
            (request.name, request.content, request.template_id, now, now, owner_user_id, owner_team_id, owner_seat_id)
        )
        await db.commit()
        novel_id = cursor.lastrowid
    finally:
        await db.close()

    async def _cleanup_empty_novel():
        # 切集失败/异常时清掉刚建的空 novel,避免库里留脏数据(FK CASCADE 自动清 chapters/scripts)
        try:
            _db = await get_db()
            try:
                cur = await _db.execute("SELECT COUNT(*) FROM chapters WHERE novel_id = ?", (novel_id,))
                has_ch = (await cur.fetchone())[0]
                if not has_ch:
                    await _db.execute("DELETE FROM novels WHERE id = ?", (novel_id,))
                    await _db.commit()
                    logger.info(f"[SCRIPT2SCRIPT-v1] 清理空 novel_id={novel_id}(无章节)")
            finally:
                await _db.close()
        except Exception as _ce:
            logger.warning(f"[SCRIPT2SCRIPT-v1] 清理空 novel_id={novel_id} 失败: {_ce}")

    try:
        # 2. 按集切分(代码定边界,不让 LLM 决定)
        chapter_count = await NovelService.parse_chapters(novel_id)
        logger.info(f"[SCRIPT2SCRIPT-v1] novel_id={novel_id} name={request.name!r} 切出 {chapter_count} 集")
        if not chapter_count:
            await _cleanup_empty_novel()
            raise HTTPException(status_code=400, detail="未能从剧本中切出有效集/章,请检查是否有「第X集」等标号")

        # 3. 逐集格式化成标准剧本(自动全量),失败明细原样回传
        #    inject_prev_context=False:逐集忠实格式化,不让上一集结尾场景串味
        convert = await ScriptService.convert_all_chapters(
            novel_id=novel_id,
            template_id=request.template_id,
            llm_config_id=request.llm_config_id,
            inject_prev_context=False,
        )
        failed = [r for r in convert.get("results", []) if not r.get("success")]
        logger.info(
            f"[SCRIPT2SCRIPT-v1] novel_id={novel_id} 格式化完成 "
            f"成功={convert.get('success_count')} 失败={convert.get('failed_count')}"
        )

        return {
            "novel_id": novel_id,
            "chapter_count": chapter_count,
            "formatted_count": convert.get("success_count", 0),
            "failed": failed,
            "message": f"剧本转剧本成功:切 {chapter_count} 集,格式化成功 {convert.get('success_count', 0)} 集"
                       + (f",失败 {len(failed)} 集" if failed else ""),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"剧本转剧本失败: {e}")
        await _cleanup_empty_novel()
        raise HTTPException(status_code=500, detail=f"剧本转剧本失败: {str(e)}")


# ============ 批量洗稿端点 ============

@router.post("/{novel_id}/batch-rewrite")
async def batch_rewrite(novel_id: int, request: BatchRewriteRequest):
    """批量洗稿"""
    # 冲突检查
    db = await get_db()
    try:
        await check_novel_running(novel_id, db)
    finally:
        await db.close()

    # 检查小说是否存在
    novel = await NovelService.get_by_id(novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="小说不存在")

    # 初始化任务状态
    _rewrite_tasks[novel_id] = {
        "status": "rewriting",
        "total": len(request.chapter_ids),
        "completed": 0,
        "failed": 0,
        "current_chapter": ""
    }

    async def _do_rewrite():
        template = await get_template_by_id(request.template_id)
        if not template:
            _rewrite_tasks[novel_id] = {"status": "failed", "total": len(request.chapter_ids), "completed": 0, "failed": len(request.chapter_ids), "current_chapter": ""}
            return

        template_content = template["content"]

        for chapter_id in request.chapter_ids:
            try:
                db = await get_db()
                try:
                    cursor = await db.execute(
                        "SELECT id, title, content FROM chapters WHERE id = ? AND novel_id = ?",
                        (chapter_id, novel_id)
                    )
                    chapter = await cursor.fetchone()
                finally:
                    await db.close()

                if not chapter:
                    _rewrite_tasks[novel_id]["failed"] += 1
                    continue

                chapter_title = chapter["title"]
                chapter_content = chapter["content"]
                _rewrite_tasks[novel_id]["current_chapter"] = chapter_title

                # 构建 prompt - 兼容多种变量名
                prompt = template_content
                replaced = False
                variable_names = ["content", "novel_content", "chapter_content", "text", "script_content"]
                for var_name in variable_names:
                    for fmt in ["{" + var_name + "}", "{{" + var_name + "}}"]:
                        if fmt in prompt:
                            prompt = prompt.replace(fmt, chapter_content)
                            replaced = True
                # 标题变量替换
                prompt = prompt.replace("{title}", chapter_title)
                prompt = prompt.replace("{{title}}", chapter_title)
                # 无占位符回退：追加到模板末尾
                if not replaced:
                    prompt = template_content + "\n\n" + chapter_content

                messages = [
                    {"role": "system", "content": "你是一位专业的小说编辑，擅长改写和润色小说内容。请对提供的章节进行洗稿改写，保持故事核心不变但用全新的表达方式。"},
                    {"role": "user", "content": prompt}
                ]

                rewritten = await LLMService.call_llm(
                    config_id=request.llm_config_id,
                    messages=messages,
                    timeout=600,
                    max_tokens=16384,
                    task_type="novel_rewrite",
                    novel_id=novel_id,
                    chapter_title=chapter_title
                )

                if rewritten and rewritten.strip():
                    rewritten_text = rewritten.strip()
                    original_len = len(chapter["content"]) if chapter["content"] else 0
                    rewritten_len = len(rewritten_text)

                    # 字数保护：如果洗稿结果不足原文的50%，视为异常，不覆盖
                    if original_len > 0 and rewritten_len < original_len * 0.5:
                        logger.warning(
                            f"洗稿结果异常：章节「{chapter_title}」原文{original_len}字，"
                            f"洗稿仅{rewritten_len}字（不足50%），跳过覆盖"
                        )
                        _rewrite_tasks[novel_id]["failed"] += 1
                        continue

                    db = await get_db()
                    try:
                        await db.execute(
                            "UPDATE chapters SET content = ? WHERE id = ?",
                            (rewritten_text, chapter_id)
                        )
                        await db.commit()
                    finally:
                        await db.close()
                    _rewrite_tasks[novel_id]["completed"] += 1
                else:
                    _rewrite_tasks[novel_id]["failed"] += 1

            except Exception as e:
                logger.error(f"洗稿章节 {chapter_id} 失败: {e}")
                _rewrite_tasks[novel_id]["failed"] += 1

        # 完成
        task = _rewrite_tasks[novel_id]
        task["current_chapter"] = ""
        if task["failed"] > 0 and task["completed"] == 0:
            task["status"] = "failed"
        elif task["failed"] > 0:
            task["status"] = "failed"
        else:
            task["status"] = "success"

    asyncio.create_task(_do_rewrite())
    return {"message": "洗稿任务已提交", "total": len(request.chapter_ids)}


@router.get("/{novel_id}/rewrite-status")
async def get_rewrite_status(novel_id: int):
    """获取洗稿进度"""
    status = _rewrite_tasks.get(novel_id, {
        "status": "idle",
        "total": 0,
        "completed": 0,
        "failed": 0,
        "current_chapter": ""
    })
    return status
