from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
IMG_DIR = ROOT / "docs" / "assets" / "user-guide-real" / "annotated"
OUT = ROOT / "docs" / "万山漫剧图文版教程.docx"


ACCENT = RGBColor(46, 116, 181)
INK = RGBColor(20, 35, 60)
MUTED = RGBColor(88, 99, 120)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8DEE9", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_east_asia(run, font_name="微软雅黑"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name="微软雅黑", size=11, color=None, bold=None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text)
    return p


def add_note_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.0)
    hdr = table.rows[0].cells
    hdr[0].text = "环节"
    hdr[1].text = "作用"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_east_asia(r)
                r.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = INK
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for idx, cell in enumerate(cells):
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_east_asia(r)
                    r.font.size = Pt(10)
                    if idx == 0:
                        r.bold = True
                        r.font.color.rgb = ACCENT
    return table


def add_image_section(doc, index, heading, intro, bullets, image_name):
    if index > 1:
        doc.add_page_break()
    h = doc.add_heading(f"{index}. {heading}", level=1)
    h.runs[0].font.color.rgb = ACCENT
    set_east_asia(h.runs[0])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, intro, color=INK)
    for b in bullets:
        add_bullet(doc, b)
    img_path = IMG_DIR / image_name
    pic = doc.add_picture(str(img_path), width=Inches(6.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    add_run(cap, f"图 {index}：{heading}", color=MUTED, size=9.5)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    set_style_font(styles["Normal"], size=10.8, color=INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.2
    set_style_font(styles["Heading 1"], size=16, color=ACCENT, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)
    set_style_font(styles["Heading 2"], size=13, color=ACCENT, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    set_style_font(styles["List Bullet"], size=10.5, color=INK)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.28)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.14)
    styles["List Bullet"].paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    add_run(title, "万山漫剧图文版教程", bold=True, color=ACCENT, size=24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    add_run(subtitle, "从小说导入到人物音频绑定、分镜生成、批量出视频", color=MUTED, size=11.5)

    p = doc.add_paragraph()
    add_run(
        p,
        "这份教程使用真实业务数据截图制作，适合发给测试用户、合作方或内部人员快速上手。"
        "当前版本的音频稳定链路是：导入人物音频样本作为音色参考，再由支持音频的视频模型在生成时使用；"
        "它不是完整的文字自动 TTS 配音系统。",
        color=INK,
    )

    doc.add_heading("整体工作链路", level=1)
    add_note_table(
        doc,
        [
            ("小说导入", "导入 txt 小说，检查章节、封面和标签。"),
            ("剧本转换", "选择章节、剧本模板和语言模型，把小说转成短剧剧本。"),
            ("信息提取", "从剧本提取人物、场景、道具，并维护资产。"),
            ("资产生成", "生成人物宫格图、场景图、道具图，也可导入成品图。"),
            ("音频绑定", "给人物或马甲导入 2-15 秒干净人声，作为音色参考。"),
            ("分镜管理", "按分镜模板生成镜头小节、人物状态和成片提示词。"),
            ("视频生成", "提交分镜、图片资产和音频参考，批量生成视频。"),
        ],
    )

    sections = [
        (
            "小说导入：项目入口",
            "导入小说后先补标签，再进入剧本、分镜和视频流程。",
            [
                "有标签的小说更适合进入后续流程。",
                "缺少屏幕模式或视觉标签时，需要点击“小说标签”补齐。",
                "常用操作包括查看章节、增量导入、导出、封面下载。",
            ],
            "01-novel-list-guide.png",
        ),
        (
            "剧本转换：小说一键转短剧",
            "选择小说、章节、剧本模板和语言模型，把小说章节转成可拍摄、可拆分镜的剧本。",
            [
                "左侧选择小说和章节，已转换章节会显示“已转换”。",
                "剧本模板决定叙事风格，语言模型负责生成文本。",
                "右侧生成结果需要检查，确认可用后保存。",
            ],
            "02-script-convert-guide.png",
        ),
        (
            "信息提取：维护人物、场景、道具资产",
            "从已保存剧本里提取人物、场景和道具。这里是视频一致性的关键页面。",
            [
                "如果提取结果是 0 个，通常说明剧本没有生成或没有保存。",
                "人物卡片里可以上传参考图、导入成品图、生成宫格图、导入音频。",
                "导入音频用于绑定角色音色，不是自动 TTS。",
            ],
            "03-extraction-assets-guide.png",
        ),
        (
            "分镜管理：剧本转镜头小节",
            "选择小说、章节、分镜模板、语言模型和风格提示词，把剧本拆成一个个视频小节。",
            [
                "分镜模板决定镜头密度、格式和风格。",
                "左侧是原始剧本，用来检查分镜是否忠实原文。",
                "右侧分镜小节会被视频生成页面读取。",
            ],
            "04-storyboard-guide.png",
        ),
        (
            "即梦视频生成：分镜批量出片",
            "选择视频通道、模型、小说章节和分镜小节，批量生成视频。",
            [
                "可选通道包括即梦 CLI、火山方舟 API、中转和星链云。",
                "不同通道支持的图片、音频、时长、比例能力不完全一样。",
                "生成日志很重要，登录失败、模型缺失、素材缺失都会在这里提示。",
            ],
            "05-video-guide.png",
        ),
        (
            "人物宫格图：保证角色一致性",
            "人物宫格图用于固定角色外观，不只是普通头像。",
            [
                "主要人物建议先生成宫格图，再批量跑视频。",
                "宫格图能降低人物跑脸、服装漂移和表情不一致。",
                "如果已有稳定角色图，也可以直接导入成品图。",
            ],
            "06-character-grid-guide.png",
        ),
    ]

    for idx, item in enumerate(sections, start=1):
        add_image_section(doc, idx, *item)

    doc.add_page_break()
    doc.add_heading("音频使用说明", level=1)
    p = doc.add_paragraph()
    add_run(p, "当前音频入口：", bold=True, color=ACCENT)
    add_run(p, "信息提取 → 人物 → 人物卡片 → 导入音频 / 麦克风按钮。")
    for item in [
        "支持 mp3、wav、m4a、ogg、flac。",
        "推荐 2 到 15 秒单人干净人声。",
        "不要背景音乐，不要多人同时说话，不要太长对白。",
        "导入后会绑定到当前人物或人物马甲。",
        "视频生成时，如果当前通道支持音频参考，系统会把人物音频作为角色音色参考一起提交。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("新用户最快上手", level=1)
    for item in [
        "设置里配置语言、图片、视频模型。",
        "导入小说并补齐小说标签。",
        "到“剧本转换”把章节转成剧本并保存。",
        "到“信息提取”提取人物、场景、道具。",
        "给主要人物生成宫格图，给需要说话的角色导入音频。",
        "到“分镜管理”生成分镜。",
        "到“即梦视频生成”批量出视频。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("后续建议补充截图", level=1)
    for item in [
        "设置页：语言、图片、视频模型配置示例。",
        "导入音频弹窗：选择本地音频文件的界面。",
        "成功生成视频后的结果页：展示最终视频、日志和素材文件夹。",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

